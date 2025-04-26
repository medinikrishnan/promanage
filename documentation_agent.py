import os
import re
import subprocess
import time
import openai
from docx import Document
from docx.shared import Inches

# ——— CONFIG ———
openai.api_key = os.getenv("OPENAI_API_KEY")
PROJECT_ROOT = os.path.dirname(__file__)
OUTPUT_MD   = os.path.join(PROJECT_ROOT, "DOCUMENTATION.md")
OUTPUT_DOCX = os.path.join(PROJECT_ROOT, "DOCUMENTATION.docx")

# Path to mmdc executable from npm installation (Add .cmd extension for Windows)
MMDC_PATH = "C:/Users/chand/AppData/Roaming/npm/mmdc.cmd"  # Full path with .cmd extension for Windows

# ——— 1) COLLECT & CHUNK ———
SKIP_DIRS = {"node_modules", ".git", "__pycache__"}
EXTS = (".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".css")

def collect_file_snippets(base_dir):
    snippets = []
    for root, dirs, files in os.walk(base_dir):
        if any(sd in root for sd in SKIP_DIRS):
            continue
        for fn in files:
            if fn.endswith(EXTS):
                path = os.path.join(root, fn)
                rel  = os.path.relpath(path, base_dir)
                try:
                    text = open(path, encoding="utf-8").read()
                except Exception as e:
                    print(f"Error reading {path}: {e}")
                    continue
                snippet = f"\n\n# File: {rel}\n```text\n{text}\n```"
                snippets.append(snippet)
    return snippets

def split_long_snippet(snip, max_chars):
    """Break a single snippet into multiple pieces of at most max_chars."""
    parts = []
    for i in range(0, len(snip), max_chars):
        parts.append(snip[i : i + max_chars])
    return parts

def chunk_snippets(snippets, max_chars=2000):
    """
    Group file‑snippets into chunks under max_chars each.
    Also split any individual snippet longer than max_chars.
    """
    chunks = []
    current = ""
    for snip in snippets:
        # if snippet itself is too big, split it first:
        if len(snip) > max_chars:
            for part in split_long_snippet(snip, max_chars):
                # push current if nonempty
                if current:
                    chunks.append(current)
                    current = ""
                chunks.append(part)
        else:
            if current and len(current) + len(snip) > max_chars:
                chunks.append(current)
                current = snip
            else:
                current += snip
    if current:
        chunks.append(current)
    return chunks

# ——— 2) PROMPTS ———
SYSTEM_PROMPT_INITIAL = """\
You are an expert software documentation generator.
Given the first slice of the codebase, produce an initial documentation in Markdown.
Include:
- Project Overview
- Installation / Setup
- Architecture Summary
- Use‑Case and Activity diagrams in Mermaid
- API references as needed

Return only Markdown."""
SYSTEM_PROMPT_EXTEND = """\
You are continuing the project documentation. Append new sections for this next slice of code.
Do NOT repeat what you already wrote; only add headings/subsections for newly seen files/features.
Maintain Markdown format, and include Mermaid diagrams where appropriate."""
USER_PROMPT = "Here is the next slice of code:\n\n{}"

# ——— 3) CHAT LOOP ———
def generate_incremental_markdown(chunks, max_retries=3):
    md = ""
    for i, chunk in enumerate(chunks):
        retries = 0
        while retries < max_retries:
            try:
                print(f"→ Processing chunk {i+1}/{len(chunks)}…")
                sys_p = SYSTEM_PROMPT_INITIAL if i == 0 else SYSTEM_PROMPT_EXTEND
                resp = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role":"system", "content": sys_p},
                              {"role":"user", "content": USER_PROMPT.format(chunk)}],
                    max_tokens=400,
                    temperature=0.2,
                    timeout=60  # Set a timeout for the API request
                )
                part = resp.choices[0].message.content
                md += "\n\n" + part
                break  # Exit the retry loop on success
            except (openai.error.APIError, openai.error.APIConnectionError) as e:
                retries += 1
                print(f"API error: {e}. Retrying {retries}/{max_retries}...")
                time.sleep(5)  # Wait before retrying
            except Exception as e:
                print(f"Unexpected error: {e}")
                break  # Exit the loop on other errors
        time.sleep(1.5)  # Throttle the requests to avoid rate-limiting
    return md

# ——— 4) SAVE MD ———
def save_markdown(md, path):
    with open(path, "w", encoding="utf-8") as f:
        f.write(md)

# ——— 5) RENDER MERMAID & BUILD DOCX ———
MERMAID_RE = re.compile(r"```mermaid\s+([\s\S]+?)```", re.MULTILINE)

def render_and_embed_mermaid(md_text, doc):
    parts = MERMAID_RE.split(md_text)
    diagram_count = 1
    for idx, chunk in enumerate(parts):
        if idx % 2 == 0:
            for block in chunk.split("\n\n"):
                txt = block.strip()
                if not txt:
                    continue
                if txt.startswith("# "):
                    doc.add_heading(txt[2:], level=1)
                elif txt.startswith("## "):
                    doc.add_heading(txt[3:], level=2)
                else:
                    doc.add_paragraph(txt)
        else:
            # this is a Mermaid block
            mmd = os.path.join(PROJECT_ROOT, f"__tmp_{idx}.mmd")  # Absolute path
            png = os.path.join(PROJECT_ROOT, f"__diag_{idx}.png")  # Absolute path
            try:
                with open(mmd, "w", encoding="utf-8") as wf:
                    wf.write(chunk.strip())

                # Run mermaid CLI to generate the PNG
                result = subprocess.run(
                    [MMDC_PATH, "-i", mmd, "-o", png, "-b", "transparent"],
                    check=False,  # Don't raise an exception on failure
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE
                )

                # Check if there were errors during diagram generation
                if result.returncode != 0:
                    print(f"Error generating diagram {diagram_count}: {result.stderr.decode()}")
                    continue  # Skip this diagram if rendering fails
                
                # Ensure the PNG file is generated before adding it to the DOCX
                if os.path.exists(png):
                    # Add the diagram to the DOCX
                    doc.add_paragraph()
                    doc.add_picture(png, width=Inches(6))
                    doc.add_paragraph(f"Figure {diagram_count}: Diagram")
                else:
                    print(f"PNG file for diagram {diagram_count} was not generated.")

            except subprocess.CalledProcessError as e:
                print(f"Error rendering Mermaid diagram: {e}")
            except Exception as e:
                print(f"Unexpected error: {e}")
            finally:
                # Clean up the temporary files
                if os.path.exists(mmd):
                    os.remove(mmd)
                if os.path.exists(png):
                    os.remove(png)
            diagram_count += 1

def build_docx(md_path, docx_path):
    md = open(md_path, encoding="utf-8").read()
    doc = Document()
    doc.add_heading("Project Documentation", 0)
    render_and_embed_mermaid(md, doc)
    doc.save(docx_path)

# ——— MAIN ———
if __name__ == "__main__":
    print(" Collecting source snippets…")
    snippets = collect_file_snippets(PROJECT_ROOT)
    chunks   = chunk_snippets(snippets)

    print(" Generating documentation in Markdown…")
    markdown = generate_incremental_markdown(chunks)

    print(f" Writing Markdown to {OUTPUT_MD}")
    save_markdown(markdown, OUTPUT_MD)

    print(" Converting to DOCX with diagrams…")
    build_docx(OUTPUT_MD, OUTPUT_DOCX)

    print(f" All done! See {OUTPUT_DOCX}")
